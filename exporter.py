import os
import json
from docx.shared import RGBColor
from openpyxl import load_workbook
from docx import Document
from docx.shared import Inches
from docx.shared import Pt

working_dir = os.getcwd()
output_txt = ''
dictname = ["oxford.json","longman.json"]

print("Which dictionary do you want to export?")
print("1.oxford\n2.longman")
dictnumber = int(input()) - 1

# 读取json文件
json_path = working_dir + '/dictionary_crawler/dictionary_crawler/spiders/' + dictname[dictnumber]
print("\nstart reading",json_path,"\n")
with open(json_path, encoding='UTF-8') as file:
    contents = file.read()

word_data = json.loads(contents)
# word_data 示例
'''
o = [
    {'displace':
         {'transitive verb':
              ['Take over the place, position, or role of (someone or something)',
               'Cause (something) to move from its proper or usual place.',
               'Force (someone) to leave their home, typically because of war, persecution, or natural disaster.',
               'Remove (someone) from a job or position of authority against their will.']
          }
     },
    {'distract': {'transitive verb': ['Prevent (someone) from giving full attention to something.',
                                      'Divert (attention) from something.',
                                      "Divert one's attention from something worrying or unpleasant by doing something different or more pleasurable.",
                                      'Perplex and bewilder.']}}, {'district': {'noun': [
        'An area of a country or city, especially one regarded as a distinct unit because of a particular characteristic.',
        'A region defined for an administrative purpose.', 'The District of Columbia; Washington, DC.'],
                                                                                'transitive verb': [
                                                                                    'Divide into districts.']}}
    , {'disposition': {'noun': ["A person's inherent qualities of mind and character.", 'An inclination or tendency.',
                                'The way in which something is placed or arranged, especially in relation to other things.',
                                'The action of arranging or ordering people or things in a particular way.',
                                'Military preparations, in particular the stationing of troops ready for attack or defense.',
                                'The action of distributing or transferring property or money to someone, in particular by bequest.',
                                'The power to deal with something as one pleases.',
                                'The determination of events by divine power.']}}, {'disposal': {
        'noun': ['The action or process of throwing away or getting rid of something.',
                 'An electrically operated device fitted to the waste pipe of a kitchen sink for grinding up food waste.',
                 'The sale of shares, property, or other assets.', 'The arrangement or positioning of something.']}}, {
        'distort': {'transitive verb': ['Pull or twist out of shape.', 'Become twisted out of shape.',
                                        'Give a misleading or false account or impression of.',
                                        'Change the form of (an electrical signal or sound wave) during transmission, amplification, or other processing.']}},
    {'disrupt': {'transitive verb': ['Interrupt (an event, activity, or process) by causing a disturbance or problem.',
                                     'Drastically alter or destroy the structure of.',
                                     '(of a company or form of technology) cause radical change in (an industry or market) by means of innovation.']}},
    {'distill': {'transitive verb': [
        'Purify (a liquid) by vaporizing it, then condensing it by cooling the vapor, and collecting the resulting liquid.',
        'Make (something, especially liquor or an essence) by distilling.',
        'Extract the essence of (something) by heating it with a solvent.',
        'Remove (a volatile constituent) of a mixture by using heat.', 'Emanate as a vapor or in minute drops.',
        'Extract the essential meaning or most important aspects of.']}}, {
        'disregard': {'transitive verb': ['Pay no attention to; ignore.'],
                      'noun': ['The action or state of disregarding or ignoring something.']}}, {'divert': {
        'transitive verb': ['Cause (someone or something) to change course or turn from one direction to another.',
                            '(of a vehicle or person) change course.',
                            'Reallocate (something, especially money or resources) to a different purpose.',
                            'Draw (the attention) of someone from something.', 'Entertain or amuse.']}}, {'diversion': {
        'noun': ['An instance of turning something aside from its course.',
                 'An alternative route for use by traffic when the usual road is temporarily closed; a detour.',
                 'An activity that diverts the mind from tedious or serious concerns; a recreation or pastime.',
                 "Something intended to distract someone's attention from something more important."]}}, {'dissipate': {
        'verb': ['(with reference to a feeling or other intangible thing) disappear or cause to disappear.',
                 'Disperse or scatter.', 'Squander or fritter away (money, energy, or resources)',
                 'Cause (energy) to be lost, typically by converting it to heat.']}}, {'divine': {
        'adjective': ['Of, from, or like God or a god.', 'Devoted to God; sacred.', 'Excellent; delightful.'],
        'noun': ['A cleric or theologian.', 'Providence or God.'],
        'transitive verb': ['Discover (something) by guesswork or intuition.',
                            'Have supernatural or magical insight into (future events)',
                            'Discover (water) by dowsing.']}}, {'disturbance': {
        'noun': ['The interruption of a settled and peaceful condition.',
                 'A breakdown of peaceful and law-abiding behavior; a riot.', 'The disruption of healthy functioning.',
                 'A local variation from normal or average wind conditions, usually a small tornado or cyclone.',
                 'Interference with rights or property; molestation.']}}, {'dizzy': {
        'adjective': ["Having or involving a sensation of spinning around and losing one's balance.",
                      "Causing a disorienting sensation of spinning around and losing one's balance.",
                      'Silly or scatterbrained (typically used of a woman)'],
        'transitive verb': ['Make (someone) feel unsteady, confused, or amazed.']}}, {'divorce': {
        'noun': ['The legal dissolution of a marriage by a court or other competent body.',
                 'A legal decree dissolving a marriage.',
                 'A separation between things which were or ought to be connected.'],
        'transitive verb': ["Legally dissolve one's marriage with (someone)",
                            'Separate or dissociate (something) from something else.',
                            'Distance or dissociate oneself from (something)']}}, {'doctrine': {
        'noun': ['A belief or set of beliefs held and taught by a Church, political party, or other group.',
                 'A stated principle of government policy, mainly in foreign or military affairs.']}}, {'dome': {
        'noun': ['A rounded vault forming the roof of a building or structure, typically with a circular base.',
                 'The revolving openable hemispherical roof of an observatory.', 'A sports stadium with a domed roof.',
                 'A thing shaped like a dome.', 'The rounded summit of a hill or mountain.',
                 'A natural vault or canopy, such as that of the sky or trees.',
                 'A rounded uplifted landform or underground structure.', 'The top of the head.',
                 'A stately building.']}}, {'dock': {'noun': [
        'A structure extending alongshore or out from the shore into a body of water, to which boats may be moored.',
        'An enclosed area of water in a port for the loading, unloading, and repair of ships.',
        'A group of enclosed areas of water along with the wharves and buildings near them.',
        'A platform for loading or unloading trucks or freight trains.',
        'A device in which a laptop computer, smartphone, or other mobile device may be placed for charging, providing access to a power supply and to peripheral devices or auxiliary features; a docking station.',
        "The solid bony or fleshy part of an animal's tail, excluding the hair.",
        'The stump left after a tail has been docked.',
        'The enclosure in a criminal court where a defendant is placed.',
        'A coarse weed of temperate regions, with inconspicuous greenish or reddish flowers. The leaves are popularly used to relieve nettle stings.'],
                                                     'intransitive verb': [
                                                         '(of a ship) tie up at a dock, especially in order to load or unload passengers or cargo.',
                                                         'Bring (a ship or boat) into a dock.',
                                                         '(of a spacecraft) join with a space station or another spacecraft in space.',
                                                         'Attach (a piece of equipment) to another.'],
                                                     'transitive verb': [
                                                         'Deduct (something, especially an amount of money)',
                                                         "Cut short (an animal's tail)"]}}, {'doom': {
        'noun': ['Death, destruction, or some other terrible fate.', '(in Christian belief) the Last Judgment.'],
        'transitive verb': ['Condemn to certain destruction or death.',
                            'Cause to have an unfortunate and inescapable outcome.']}}, {'dramatic': {
        'adjective': ['Relating to drama or the performance or study of drama.',
                      '(of an event or circumstance) sudden and striking.', 'Exciting or impressive.',
                      '(of a person or their behavior) intending or intended to create an effect; theatrical.']}}, {
        'drag': {'verb': ['Pull (someone or something) along forcefully, roughly, or with difficulty.',
                          'Take (someone) to or from a place or event, despite their reluctance.',
                          'Go somewhere wearily, reluctantly, or with difficulty.',
                          "(of a person's clothes or an animal's tail) trail along the ground.",
                          'Catch hold of and pull (something)',
                          '(of a ship) trail (an anchor) along the seabed, causing the ship to drift.',
                          '(of an anchor) fail to hold, causing a ship or boat to drift.',
                          'Search the bottom of (a river, lake, or the sea) with grapnels or nets.',
                          '(of time, events, or activities) pass slowly and tediously.',
                          'Move (an icon or other image) across a computer screen using a tool such as a mouse.',
                          'Engage in a drag race.'],
                 'noun': ['The action of pulling something forcefully or with difficulty.',
                          'The longitudinal retarding force exerted by air or other fluid surrounding a moving object.',
                          'A person or thing that impedes progress or development.',
                          'Unnatural motion of a fishing fly caused by the pull of the line.',
                          'An iron shoe that can be applied as a brake to the wheel of a cart or wagon.',
                          'A boring or tiresome person or thing.', 'An act of inhaling smoke from a cigarette.',
                          "Clothing more conventionally worn by the opposite sex, especially women's clothes worn by a man.",
                          'A street or road.', 'A thing that is pulled along the ground or through water.',
                          'A harrow used for breaking up the surface of land.',
                          'An apparatus for dredging a river or for recovering the bodies of drowned people from a river, a lake, or the sea.',
                          'A strong-smelling lure drawn before hounds as a substitute for a fox or other hunted animal.',
                          'A hunt using a strong-smelling lure.', 'Influence over other people.',
                          'One of the basic patterns (rudiments) of drumming, consisting of a stroke preceded by two grace notes, which are usually played with the other stick.',
                          'A private vehicle like a stagecoach, drawn by four horses.']}}, {'donate': {
        'transitive verb': ['Give (money or goods) for a good cause, for example to a charity.',
                            "Allow the removal of (blood or an organ) from one's body for transplantation, transfusion, or other use."]}},
    {'drainage': {'noun': ['The action or process of draining something.',
                           'The means of removing surplus water or liquid waste; a system of drains.']}}, {'domain': {
        'noun': ['An area of territory owned or controlled by a ruler or government.',
                 'A specified sphere of activity or knowledge.',
                 'A discrete region of magnetism in ferromagnetic material.',
                 'A distinct subset of the internet with addresses sharing a common suffix or under the control of a particular organization or individual.',
                 'The set of possible values of the independent variable or variables of a function.']}},
    {'drastic': {'adjective': ['Likely to have a strong or far-reaching effect; radical and extreme.']}}, {'drawback': {
        'noun': ['A feature that renders something less acceptable; a disadvantage or problem.',
                 'An amount of excise or import duty remitted on imported goods that the importer re-exports rather than sells domestically.']}},
    {'dread': {
        'transitive verb': ['Anticipate with great apprehension or fear.', 'Regard with great awe or reverence.'],
        'noun': ['Great fear or apprehension.', 'A person with dreadlocks.', 'Dreadlocks.'],
        'adjective': ['Greatly feared; dreadful.', 'Regarded with awe; greatly revered.']}}, {'dual': {
        'adjective': ['Consisting of two parts, elements, or aspects.',
                      '(in some languages) denoting an inflection that refers to exactly two people or things (as distinct from singular and plural)',
                      '(in an aircraft) using dual controls.',
                      '(of a theorem, expression, etc.) related to another by the interchange of particular pairs of terms, such as “point” and “line.”.'],
        'noun': ['A dual form of a word.', 'The dual number.',
                 'A theorem, expression, etc., that is dual to another.']}}, {'dubious': {
        'adjective': ['Hesitating or doubting.', 'Not to be relied upon; suspect.', 'Morally suspect.',
                      'Of questionable value.']}}, {'drought': {
        'noun': ['A prolonged period of abnormally low rainfall, leading to a shortage of water.',
                 'A prolonged absence of a specified thing.', 'Thirst.']}}, {
        'dwell': {'verb': ['Live in or at a specified place.'],
                  'noun': ['A slight regular pause in the motion of a machine.']}},
    {'duration': {'noun': ['The time during which something continues.']}}, {'drill': {'noun': [
        'A hand tool, power tool, or machine with a rotating cutting tip or reciprocating hammer or chisel, used for making holes.',
        'A tool used by a dentist for cutting away part of a tooth before filling it.',
        'Instruction or training in military exercises.',
        'Intensive instruction or training in something, typically by means of repeated exercises.',
        'A rehearsal of the procedure to be followed in an emergency.',
        'The correct or recognized procedure or way of doing something.',
        'A predatory mollusk that bores into the shells of other mollusks in order to feed on the soft tissue.',
        'A machine that makes small furrows, sows seed in them, and then covers the sown seed.',
        'A small furrow, especially one made by a drill.', 'A ridge with a furrow on top made by a drill.',
        'A row of plants sown in a furrow made by a drill.',
        'A dark brown baboon with a short tail and a naked blue or purple rump, found in the rainforests of West Africa.',
        'A coarse twilled cotton or linen fabric.'], 'transitive verb': [
        'Produce (a hole) in something by or as if by boring with a drill.',
        'Make a hole in (something) by boring with a drill.', 'Make a hole in or through something by using a drill.',
        'Sink a borehole in order to obtain a certain substance, typically oil or water.',
        '(of a dentist) cut away part of (a tooth) before filling it.',
        '(of a sports player) hit, throw, or kick (a ball or puck) hard and in a straight line.',
        'Subject (someone) to military training exercises.', '(of a person) take part in military training exercises.',
        'Instruct (someone) in something by the means of repeated exercises or practice.',
        '(of a person or machine) sow (seed) with a drill.', 'Plant (the ground) in furrows.']}}, {'ease': {
        'noun': ['Absence of difficulty or effort.', 'Absence of rigidity or discomfort; poise.',
                 "Freedom from worries or problems, especially about one's material situation."],
        'verb': ['Make (something unpleasant, painful, or intense) less serious or severe.',
                 'Become less serious or severe.', 'Make (something) happen more easily; facilitate.',
                 '(of share prices, interest rates, etc.) decrease in value or amount.', 'Slacken a rope or sail.',
                 'Move carefully or gradually.', 'Move (someone or something) carefully or gradually.']}}, {'ecology': {
        'noun': [
            'The branch of biology that deals with the relations of organisms to one another and to their physical surroundings.']}},
    {'eccentric': {'adjective': ['(of a person or their behavior) unconventional and slightly strange.',
                                 '(of a thing) not placed centrally or not having its axis or other part placed centrally.',
                                 '(of a circle) not centered on the same point as another.',
                                 '(of an orbit) not circular.'],
                   'noun': ['A person of unconventional and slightly strange views or behavior.',
                            'A disc or wheel mounted eccentrically on a revolving shaft in order to transform rotation into backward-and-forward motion, e.g. a cam in an internal combustion engine.']}},
    {'earnings': {'plural noun': ['Money obtained in return for labor or services.',
                                  'Income derived from an investment or product.']}}, {'eclipse': {'noun': [
        'An obscuring of the light from one celestial body by the passage of another between it and the observer or between it and its source of illumination.',
        'A loss of significance, power, or prominence in relation to another person or thing.',
        'A phase during which the distinctive markings of a bird (especially a male duck) are obscured by molting of the breeding plumage.'],
                                                                                                   'transitive verb': [
                                                                                                       '(of a celestial body) obscure the light from or to (another celestial body)',
                                                                                                       'Deprive (someone or something) of significance, power, or prominence.',
                                                                                                       'Obscure or block out (light)']}}]
'''

# 创建word文档
print("start writing:")
document = Document()

for i in word_data:
    # 遍历每个单词
    # 写入  para指每个段落  一个段落一个单词

    """
    example：
    word1
        noun：meaning
             example：
        verb：meaning
             example
    word2：
    …………  
    """
    para = document.add_paragraph()
    run = para.add_run(list(i.keys())[0])
    font = run.font

    font.name = 'Calibri'
    font.size = Pt(14)
    # para = document.add_paragraph(list(i.keys())[0],)  #单词本身
    # para.add_run(list(i.keys())[0])
    # list(i.keys())[0] 是单词本身的值（dict返回列表是假的，要转换）

    para.add_run('\n')  # 换行

    for characteristic in i.get(list(i.keys())[0]).items():
        para.add_run('    ')

        # 词性
        print(characteristic)
        run = para.add_run(characteristic[0])
        font = run.font
        font.name = '等线'
        font.size = Pt(8)
        font.color.rgb = RGBColor(250,200,200)
        font.italic = True
        # para.add_run(characteristic[0])  #词性

        para.add_run('\n')
        para.add_run('      ')

        # 释义
        seq = 0
        for definitions in characteristic[1]:
            seq += 1
            para.add_run("\t"+str(seq)+". ")
            run = para.add_run(definitions)
            font = run.font
            font.name = 'Consolas'
            font.size = Pt(8)

            font.italic = False
            font.bold = True
            para.add_run("\n")
        # para.add_run(characteristic[1])  #释义

document.save('demo.docx')